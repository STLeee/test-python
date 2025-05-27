from enum import StrEnum
import json
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, Field
from latex2word import LatexToWordElement


quiz_type_index_mapping = {
    1: "壹",
    2: "貳",
    3: "參",
    4: "肆",
    5: "伍",
    6: "陸",
    7: "柒",
    8: "捌",
    9: "玖",
    10: "拾",
}

option_index_mapping = {
    1: "A",
    2: "B",
    3: "C",
    4: "D",
}

level_mapping = {
    "junior": "國中",
    "elementary": "國小",
}

quiz_type_mapping = {
    "mcq": "選擇題",
    "saq": "簡答題",
}

quiz_category_mapping = {
    "bopomofo": "字音字型",
    "terms-meaning": "詞意理解",
    "reading-comprehension": "文意理解",
    "basic-concept": "基礎觀念",
    "other": "其它",
}


class QuizType(StrEnum):
    MULTIPLE_CHOICE = "mcq"
    SHORT_ANSWER = "saq"


class QuizCategory(StrEnum):
    # For Chinese
    BOPOMOFO = "bopomofo"
    TERMS_MEANING = "terms-meaning"
    READING_COMPREHENSION = "reading-comprehension"
    # For Science & Social Studies & Geography & History & Civics
    BASIC_CONCEPT = "basic-concept"
    # Default
    OTHER = "other"


class Quiz(BaseModel):
    quiz_type: QuizType = Field(default=QuizType.MULTIPLE_CHOICE)
    quiz_category: QuizCategory = Field(default=QuizCategory.OTHER)
    source: str
    question: str
    options: list[str]
    explanation: str
    answer: int


class Quizzes(BaseModel):
    academic_year: int
    level: str
    grade: str
    semester: str
    subject: str
    chapter: str
    title: str
    quizzes: list[Quiz]


def replace_quizzes_info(quizzes: Quizzes, document: DocumentObject) -> str:
    keywords = {
        "{{year}}": str(quizzes.academic_year),
        "{{level}}": level_mapping.get(quizzes.level, quizzes.level),
        "{{grade}}": quizzes.grade,
        "{{semester}}": quizzes.semester,
        "{{subject}}": quizzes.subject,
        "{{chapter}}": quizzes.chapter,
        "{{title}}": quizzes.title,
    }
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text_items = paragraph.text.split(" ")

                    paragraph.text = ""
                    for text_item in text_items:
                        for keyword, value in keywords.items():
                            text_item = text_item.replace(keyword, value)

                        paragraph.add_run(text_item + " ", style="Title_answers" if "答案卷" in text_item else None)


def add_text_with_latex(text: str, paragraph: Paragraph):
    text_parts = text.split("$")
    for part in text_parts:
        if part.startswith("\\frac"):
            latex_to_word = LatexToWordElement(part)
            latex_to_word.add_latex_to_paragraph(paragraph)
        else:
            paragraph.add_run(part)


def insert_horizontal_line(paragraph: Paragraph) -> None:
    p = paragraph._p  # Access the internal XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert(0, pBdr)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')  # Line style: single
    bottom.set(qn('w:sz'), '6')  # Line thickness
    bottom.set(qn('w:space'), '1')  # Space between text and line
    bottom.set(qn('w:color'), 'auto')  # Line color
    pBdr.append(bottom)


def add_next_paragraph(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    document: DocumentObject = paragraph._parent
    next_paragraph: Paragraph = document.add_paragraph(text=text, style=style)
    paragraph._p.addnext(next_paragraph._element)
    return next_paragraph


def add_quizzes(
    index: int, quiz_type: QuizType, quizzes: list[Quiz], paragraph: Paragraph, is_answers: bool = False
) -> Paragraph:
    paragraph.add_run(
        f"{quiz_type_index_mapping.get(index, index)}、{quiz_type_mapping.get(quiz_type, quiz_type)} (每題 ___ 分。共 ____ 分)："
    )
    paragraph.style = "quizzes_title"

    quiz_index = 0
    for quiz in quizzes:
        if quiz.quiz_type != quiz_type:
            continue

        if quiz_type == QuizType.MULTIPLE_CHOICE:
            answer_text = option_index_mapping.get(quiz.answer + 1, quiz.answer + 1) if is_answers else "  "
            paragraph = add_next_paragraph(paragraph, text="（", style="quiz_question_mcq")
            paragraph.add_run(text=answer_text, style="quiz_option_answer")
            paragraph.add_run(text=f"）{quiz_index+1}. ")
            add_text_with_latex(quiz.question, paragraph)
            for i, option in enumerate(quiz.options):
                paragraph = add_next_paragraph(paragraph, style="quiz_option")
                paragraph.add_run(text=f"（{option_index_mapping.get(i + 1, i + 1)}）")
                add_text_with_latex(option, paragraph)
            if is_answers:
                paragraph = add_next_paragraph(paragraph, style="quiz_explanation_mcq")
                paragraph.add_run(text="詳解：\r", style="quiz_explanation_title_mcq")
                add_text_with_latex(quiz.explanation, paragraph)
        elif quiz_type == QuizType.SHORT_ANSWER:
            paragraph = add_next_paragraph(
                paragraph,
                text=f"{quiz_index+1}. ",
                style="quiz_question_saq",
            )
            add_text_with_latex(quiz.question, paragraph)
            if not is_answers:
                paragraph = add_next_paragraph(paragraph)
                insert_horizontal_line(paragraph)
                paragraph = add_next_paragraph(paragraph, style="horizontal_line")
                paragraph = add_next_paragraph(paragraph)
                insert_horizontal_line(paragraph)
                paragraph = add_next_paragraph(paragraph)
            else:
                paragraph = add_next_paragraph(paragraph, style="quiz_explanation")
                paragraph.add_run(text="參考答案：\r", style="quiz_explanation_title")
                add_text_with_latex(quiz.explanation, paragraph)
                insert_horizontal_line(paragraph)

        quiz_index += 1
    return paragraph


def main():
    # load quizzes.json file
    quizzes_file = "quizzes.json"
    try:
        with open(quizzes_file, "r", encoding="utf-8") as f:
            quizzes_data = f.read()
    except FileNotFoundError:
        print(f"Error: The file {quizzes_file} does not exist.")
        return
    except Exception as e:
        print(f"Error reading {quizzes_file}: {e}")
        return
    try:
        quizzes = Quizzes(**json.loads(quizzes_data))
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON from {quizzes_file}: {e}")
        return
    except TypeError as e:
        print(f"Error creating Quiz objects: {e}")
        return

    document = Document("template.docx")

    replace_quizzes_info(quizzes, document)

    for paragraph in document.paragraphs:
        if "{{quizzes}}" in paragraph.text:
            paragraph.text = ""
            paragraph = add_quizzes(1, QuizType.MULTIPLE_CHOICE, quizzes.quizzes, paragraph)
            paragraph = add_next_paragraph(paragraph)
            paragraph = add_quizzes(1, QuizType.SHORT_ANSWER, quizzes.quizzes, paragraph)
        elif "{{answers}}" in paragraph.text:
            paragraph.text = ""
            paragraph = add_quizzes(1, QuizType.MULTIPLE_CHOICE, quizzes.quizzes, paragraph, is_answers=True)
            paragraph = add_next_paragraph(paragraph)
            paragraph = add_quizzes(2, QuizType.SHORT_ANSWER, quizzes.quizzes, paragraph, is_answers=True)

    document.save("demo.docx")


if __name__ == "__main__":
    main()
    print("Document created successfully.")
