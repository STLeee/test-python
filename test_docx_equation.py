import re
import sys

from docx import Document
from docx.text.paragraph import Paragraph
from latex2word import LatexToWordElement


def add_latex_to_paragraph(latex: str, paragraph: Paragraph):
    try:
        latex_element = LatexToWordElement(latex)
        latex_element.add_latex_to_paragraph(paragraph)
    except Exception as e:
        print(f"Error processing LaTeX: {e}")
        paragraph.add_run(latex)


def add_text_with_latex(text: str, paragraph: Paragraph):
    regex_latex = r"\$\$(.*?)\$\$|(?<!\\)\$([^\$\n]+?)(?<!\\)\$"
    matches = re.finditer(regex_latex, text, re.DOTALL)
    last_end = 0
    for match in matches:
        # Add normal text before the LaTeX
        paragraph.add_run(text[last_end : match.start()])
        # Add LaTeX content
        latex_content = match.group(1) or match.group(2)
        add_latex_to_paragraph(latex_content, paragraph)
        # Update last_end to the end of the current match
        last_end = match.end()
    # Add any remaining text after the last LaTeX
    paragraph.add_run(text[last_end:])


def main():
    text_file = "test_docx_equation.txt"

    # load text file
    try:
        with open(text_file, "r", encoding="utf-8") as f:
            text_data = f.read()
    except FileNotFoundError:
        print(f"File {text_file} not found.")
        return 1
    except Exception as e:
        print(f"Error reading {text_file}: {e}")
        return 1

    document = Document()

    try:
        for block in text_data.split("==="):
            paragraph = document.add_paragraph()
            add_text_with_latex(block.strip(), paragraph)
            document.add_paragraph("\n" + "─" * 50 + "\n")
    except Exception as e:
        print(f"Error processing text data: {e}")
        return 1

    document.save("test_docx_equation.docx")


if __name__ == "__main__":
    sys.exit(main())
    print("Document created successfully.")
